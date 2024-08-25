import openpyxl
from openpyxl.styles import Font,PatternFill,GradientFill,Alignment,Border,Side

import docx

import pandas as pd

# ---------------xlsx----------------

data = pd.read_excel('a.xls', sheet_name='Sheet1') #pd数组

# 打开和保存
book = openpyxl.load_workbook('a.xlsx')
sheet = book.sheetnames[0]
book.save('b.xlsx')

# 读取和写入
if book[sheet]['A1'].value is not None:
    print(book[sheet]['A1'].value)
book[sheet]['A1'] = 'hehe'

# 改变行列宽高
book[sheet].row_dimensions[2].height = 30
book[sheet].column_dimensions["B"].width = 30

# 改变字体
book[sheet]['A1'].font = Font(
    name="微软雅黑",   # 字体
    size=15,         # 字体大小
    color="0000FF",  # 字体颜色，用16进制rgb表示
    bold=True,       # 是否加粗，True/False
    italic=True,     # 是否斜体，True/False
    strike=None,     # 是否使用删除线，True/False
    underline=None,  # 下划线, 可选'singleAccounting', 'double', 'single', 'doubleAccounting'
)


# 改变填充
book[sheet]['A1'].fill = PatternFill(
    patternType="solid",  # 填充类型，可选none、solid、darkGray、mediumGray、lightGray、lightDown、lightGray、lightGrid
    fgColor="F562a4",  # 这个才是背景色！
    bgColor="0000ff",  # 不知道是啥
    # fill_type=None,  # 填充类型
    # start_color=None, # 前景色，16进制rgb
    # end_color=None    # 背景色，16进制rgb
)
book[sheet]['A1'].fill = GradientFill(
    degree=60,  # 角度
    stop=("000000", "FFFFFF")  # 渐变颜色，16进制rgb
)

# 改变对齐
book[sheet]['A1'].alignment = Alignment(
    horizontal='left',  # 水平对齐，可选general、left、center、right、fill、justify、centerContinuous、distributed
    vertical='top',  # 垂直对齐， 可选top、center、bottom、justify、distributed
    text_rotation=0,  # 字体旋转，0~180整数
    wrap_text=False,  # 是否自动换行
    shrink_to_fit=False,  # 是否缩小字体填充
    indent=0,  # 缩进值
)

# 改变边框样式
side = Side(
    style="medium",  # 边框样式，可选dashDot、dashDotDot、dashed、dotted、double、hair、medium、mediumDashDot、mediumDashDotDot、mediumDashed、slantDashDot、thick、thin
    color="ff66dd",  # 边框颜色，16进制rgb表示
)
book[sheet]['A1'].border = Border(
    top=side,  # 上
    bottom=side,  # 下
    left=side,  # 左
    right=side,  # 右
    diagonal=side  # 对角线
)


# ---------------docx----------------

# 打开和保存
doc = docx.Document('a.docx')

texts = [p.text for p in doc.paragraphs]

tables = doc.tables

tables[0].rows #行数
tables[0].columns #列数

tables[0].cell(0,0).text #获取值
